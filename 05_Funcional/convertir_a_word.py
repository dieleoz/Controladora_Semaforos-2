#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de entregables Word (.docx) para la carpeta 05_Funcional.

FUENTE ÚNICA DE VERDAD: los archivos .md de esta carpeta.
------------------------------------------------------------------------------
La versión anterior tenía el contenido del checklist HARDCODEADO en este script,
independiente del .md. Ambas copias divergieron: el .md se corregía y el .docx
—que es el documento que se imprime y se firma— seguía con el texto viejo.

Ahora cada .docx se construye leyendo su .md correspondiente. Editar el .md es
suficiente; volver a correr este script regenera el Word.

Uso:
    python 05_Funcional/convertir_a_word.py            # genera los 6 documentos
    python 05_Funcional/convertir_a_word.py 3 4        # solo los numerados 3 y 4

Funciona desde cualquier directorio de trabajo.
"""

import os
import re
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Rutas resueltas desde la ubicación del script, no desde el cwd.
DIR = os.path.dirname(os.path.abspath(__file__))

AZUL = RGBColor(0, 51, 102)
GRIS = RGBColor(90, 90, 90)
ROJO = RGBColor(150, 0, 0)


# ---------------------------------------------------------------- utilidades

def _sombrear(elemento, hex_fill):
    """Aplica color de fondo a una celda (_Cell) o a un párrafo (Paragraph).

    python-docx expone las propiedades de sombreado en el elemento XML subyacente
    (`_tc` para celdas, `_p` para párrafos), no en el objeto de alto nivel.
    """
    shd = docx.oxml.OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    if hasattr(elemento, "_tc"):          # _Cell
        elemento._tc.get_or_add_tcPr().append(shd)
    else:                                  # Paragraph
        elemento._p.get_or_add_pPr().append(shd)


def _texto_enriquecido(parrafo, texto, base_bold=False, color=None, size=None):
    """Vuelca `texto` en el párrafo interpretando **negrita** y `código`."""
    # Limpia enlaces markdown [etiqueta](destino) -> etiqueta
    texto = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", texto)
    # Limpia LaTeX simple usado en los manuales
    texto = texto.replace("$\\rightarrow$", "→").replace("$\\leftrightarrow$", "↔")
    texto = re.sub(r"\$([^$]*)\$", r"\1", texto)

    for trozo in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", texto):
        if not trozo:
            continue
        run = parrafo.add_run()
        if trozo.startswith("**") and trozo.endswith("**"):
            run.text = trozo[2:-2]
            run.bold = True
        elif trozo.startswith("`") and trozo.endswith("`"):
            run.text = trozo[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            run.text = trozo
            run.bold = base_bold
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)


def _fila_tabla(linea):
    """'| a | b |' -> ['a', 'b']"""
    return [c.strip() for c in linea.strip().strip("|").split("|")]


def _es_separador_tabla(linea):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", linea.strip()))


# ------------------------------------------------------------- conversión md

def markdown_a_word(ruta_md, ruta_docx):
    with open(ruta_md, "r", encoding="utf-8") as f:
        lineas = f.read().splitlines()

    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1.0)

    i = 0
    while i < len(lineas):
        linea = lineas[i]
        desnuda = linea.strip()

        # --- Bloque de código cercado -------------------------------------
        if desnuda.startswith("```"):
            i += 1
            buffer = []
            while i < len(lineas) and not lineas[i].strip().startswith("```"):
                buffer.append(lineas[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run("\n".join(buffer))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            _sombrear(p, "F4F4F4")
            continue

        # --- Tabla ---------------------------------------------------------
        if desnuda.startswith("|") and i + 1 < len(lineas) and _es_separador_tabla(lineas[i + 1]):
            cabecera = _fila_tabla(desnuda)
            i += 2
            filas = []
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                filas.append(_fila_tabla(lineas[i]))
                i += 1

            tabla = doc.add_table(rows=1, cols=len(cabecera))
            tabla.style = "Table Grid"
            for idx, celda_txt in enumerate(cabecera):
                celda = tabla.rows[0].cells[idx]
                celda.text = ""
                _texto_enriquecido(celda.paragraphs[0], celda_txt, base_bold=True, size=9.5)
                _sombrear(celda, "DCE6F1")
            for fila in filas:
                celdas = tabla.add_row().cells
                for idx, celda_txt in enumerate(fila[:len(cabecera)]):
                    celdas[idx].text = ""
                    _texto_enriquecido(celdas[idx].paragraphs[0], celda_txt, size=9.5)
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # --- Separador horizontal ------------------------------------------
        if re.fullmatch(r"-{3,}", desnuda):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # --- Encabezados ----------------------------------------------------
        m = re.match(r"^(#{1,4})\s+(.*)", desnuda)
        if m:
            nivel, texto = len(m.group(1)), m.group(2)
            p = doc.add_paragraph()
            if nivel == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14 if nivel <= 2 else 10)
            p.paragraph_format.space_after = Pt(6)
            _texto_enriquecido(p, texto, base_bold=True, color=AZUL,
                               size={1: 16, 2: 13, 3: 11.5, 4: 10.5}[nivel])
            i += 1
            continue

        # --- Cita / aviso ----------------------------------------------------
        if desnuda.startswith(">"):
            buffer = []
            while i < len(lineas) and lineas[i].strip().startswith(">"):
                buffer.append(re.sub(r"^\s*>\s?", "", lineas[i]))
                i += 1
            texto = " ".join(l.strip() for l in buffer if l.strip())
            texto = re.sub(r"^#{1,4}\s*", "", texto)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(8)
            _texto_enriquecido(p, texto, color=ROJO, size=10)
            _sombrear(p, "FFF4E5")
            continue

        # --- Casilla de verificación -----------------------------------------
        m = re.match(r"^-\s*\[\s*\]\s+(.*)", desnuda)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("☐  ")
            run.font.size = Pt(12)
            run.bold = True
            _texto_enriquecido(p, m.group(1), size=10.5)
            i += 1
            continue

        # --- Viñeta / numerada -------------------------------------------------
        m = re.match(r"^(\s*)[-*]\s+(.*)", linea)
        if m:
            sangria = 0.25 + 0.25 * (len(m.group(1)) // 2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(sangria)
            p.paragraph_format.space_after = Pt(3)
            _texto_enriquecido(p, m.group(2), size=10.5)
            i += 1
            continue

        m = re.match(r"^(\s*)\d+\.\s+(.*)", linea)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            _texto_enriquecido(p, m.group(2), size=10.5)
            i += 1
            continue

        # --- Línea en blanco / párrafo normal -----------------------------------
        if not desnuda:
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _texto_enriquecido(p, desnuda, size=10.5)
        i += 1

    # Pie con la fuente de origen, para que nadie edite el .docx a mano.
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _texto_enriquecido(
        pie,
        f"Documento generado automáticamente desde {os.path.basename(ruta_md)}. "
        "No editar este .docx directamente: los cambios se pierden al regenerarlo.",
        color=GRIS, size=8)

    doc.save(ruta_docx)
    return ruta_docx


def main():
    filtro = set(sys.argv[1:])
    generados, fallidos = [], []

    for nombre in sorted(os.listdir(DIR)):
        if not nombre.endswith(".md") or nombre.upper().startswith("README"):
            continue
        num = nombre.split("_", 1)[0]
        if filtro and num not in filtro:
            continue
        ruta_md = os.path.join(DIR, nombre)
        ruta_docx = os.path.join(DIR, nombre[:-3] + ".docx")
        try:
            markdown_a_word(ruta_md, ruta_docx)
            generados.append(os.path.basename(ruta_docx))
        except Exception as e:  # noqa: BLE001 - se reporta y se sigue con los demás
            fallidos.append(f"{nombre}: {e}")

    print("Documentos Word generados:")
    for g in generados:
        print(f"   OK  {g}")
    for f in fallidos:
        print(f"   ERROR  {f}")
    if fallidos:
        sys.exit(1)


if __name__ == "__main__":
    main()
