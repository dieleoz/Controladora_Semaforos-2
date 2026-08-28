import docx
import sys
import re

def markdown_to_docx(md_path, docx_path):
    try:
        doc = docx.Document()
    except Exception as e:
        import docx
        doc = docx.Document()

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line[2:])
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', line)
            p.add_run(text)
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)

markdown_to_docx('D:\\@Proyect\\Controladora_Semaforos\\05_Funcional\\5_Manual_Puente_ESP32.md', 'D:\\@Proyect\\Controladora_Semaforos\\05_Funcional\\5_Manual_Puente_ESP32.docx')
